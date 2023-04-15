from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from quest_answ_randomizer import Randomizer


class Generator:
    def __init__(self, xlsx_file_name: str, list_nbr_question: list, nbr_copies: int, qcm_file_prefix: str):
        for i in range(nbr_copies):
            quest_and_answer = Randomizer(xlsx_file_name, list_nbr_question).dico_quest_ans
            file = Document()
            for sheet in quest_and_answer:
                quest_and_answer_ws = quest_and_answer[sheet]
                for question in quest_and_answer_ws.keys():
                    quest_and_answer_wsiq = quest_and_answer_ws[question]
                    pg = file.add_paragraph(style='List Number')
                    pg.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    pg.paragraph_format.space_after = Pt(0)
                    pg_run = pg.add_run(question)
                    pg_run.font.size = Pt(12)
                    pg_run.font.name = 'Calibri Light'
                    for i1 in range(len(quest_and_answer_wsiq)):
                        pg = file.add_paragraph()
                        pg.paragraph_format.left_indent = Inches(0.24)
                        pg.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pg.paragraph_format.space_after = Pt(0)
                        pg_run = pg.add_run(f"□ {chr(65 + i1)}. {quest_and_answer_wsiq[i1]}")
                        pg_run.font.size = Pt(12)
                        pg_run.font.name = 'Calibri Light'
                    file.add_paragraph()
            file.save(f'{qcm_file_prefix}_{i+1}.docx')


Generator('exemple', [2, 3], 1, 'my_exemple_qcm')
