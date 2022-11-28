from csv import reader as read
from docx import Document
from random import randint
from docx.shared import Pt,Inches
class qcm_generator:
    def read_csv(csv_name):
        csv_open=open(f'{csv_name}.csv','r')
        csv_read=read(csv_open,delimiter=';')
        quest_answ={}
        for row in csv_read:
            if row[0]!='':
                quest=row[1]
                quest_answ[quest]=[row[2]]
            else:
                quest_answ[quest].append(row[2])
        csv_open.close()
        return quest_answ
    def generate(self,nbr_qcm:int,csv_name:str):
        for i_1 in range(nbr_qcm):
            quest_answ=self.read_csv(csv_name)
            qcm=Document()
            index_quest={}
            index=0
            for quest in quest_answ:
                index+=1
                index_quest[index]=quest
            for i_2 in range(len(quest_answ)):
                index_quest={}
                index=0
                for quest in quest_answ:
                    index+=1
                    index_quest[index]=quest
                random_quest_nbr=randint(1,len(index_quest))
                pg=qcm.add_paragraph(style='List Number').add_run(f'{index_quest[random_quest_nbr]}:')
                pg.font.size=Pt(12)
                pg.font.name='Calibri Light'
                
                index_answ={}
                index=0
                for answ in quest_answ[index_quest[random_quest_nbr]]:
                    index+=1
                    index_answ[index]=answ
                for i_3 in range(len(index_answ)):
                    index_answ={}
                    index=0
                    for answ in quest_answ[index_quest[random_quest_nbr]]:
                        index+=1
                        index_answ[index]=answ
                    random_answ_nbr=randint(1,len(index_answ))
                    pg=qcm.add_paragraph()
                    pg.paragraph_format.left_indent=Inches(1.97)
                    pg_run=pg.add_run(f"□ {chr(65+i_3)}. {index_answ[random_answ_nbr]}")
                    pg_run.font.size=Pt(12)
                    pg_run.font.name='Calibri Light'
                    del quest_answ[index_quest[random_quest_nbr]][random_answ_nbr-1]
                qcm.add_paragraph()
                del quest_answ[index_quest[random_quest_nbr]]
            qcm.save(f"QCM_{i_1+1}.docx")